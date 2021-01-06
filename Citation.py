


import selenium
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
import pyautogui
import time
import docx
import os
import sys

from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

#setting up the condition
r = 0
def annoyingad():
	pyautogui.moveTo(979,568)
	time.sleep(0.2)
	pyautogui.click()
	time.sleep(23.5)

time.sleep(7)

#Removing formatting
pyautogui.hotkey('ctrl', 'a')
pyautogui.hotkey('ctrl', 'x')
pyautogui.hotkey('ctrl', 'v')
time.sleep(1)
pyautogui.typewrite(["ctrl"])
pyautogui.typewrite(["t"])

#Formatting the font
pyautogui.hotkey('ctrl', 'a')
pyautogui.hotkey('alt')
pyautogui.typewrite('h')
pyautogui.typewrite('f')
pyautogui.typewrite('s')
pyautogui.hotkey('8')
pyautogui.hotkey('enter')


#Formatting the margins
pyautogui.hotkey('alt')
pyautogui.typewrite('p')
pyautogui.typewrite('m')
pyautogui.hotkey('tab')
pyautogui.hotkey('enter')

time.sleep(2)

#Copy Pasting the file
pyautogui.hotkey('win', 'd')
pyautogui.hotkey('ctrl', 'c')
pyautogui.hotkey('ctrl', 'v')


PATH= "D:\Python Projects\chromedriver.exe"
driver = webdriver.Chrome(PATH)
driver.get('https://www.citationmachine.net/apa/cite-a-website')

time.sleep(2)

pyautogui.hotkey('win', 'up')

doc = docx.Document('testing - Copy.docx')
tt = len(doc.paragraphs) - 1


#Begin the for loop
for z in range (0, tt, ):



	search = driver.find_element_by_id("search-input")

	search.send_keys(doc.paragraphs[z].text)
	search.send_keys(Keys.RETURN)


#cite
	time.sleep(7)
	pyautogui.moveTo(15,602)
	pyautogui.click()
	pyautogui.hotkey('tab')
	pyautogui.hotkey('tab')
	pyautogui.hotkey('tab')
	pyautogui.hotkey('tab')
	pyautogui.hotkey('tab')
	pyautogui.hotkey('tab')
	
	if r > 0:
		pyautogui.hotkey('tab')

	pyautogui.hotkey('enter')

#Continue
	time.sleep(3)
	pyautogui.hotkey('tab')
	if r > 0:
		pyautogui.hotkey('tab')
	pyautogui.hotkey('enter')

#complete citation
	time.sleep(1)
	pyautogui.moveTo(15,602)
	pyautogui.click()
	pyautogui.hotkey('pagedown')

	time.sleep(2)
	completebutton = driver.find_element_by_xpath('//*[@id="citation_form"]/div[2]/button')
	completebutton.click()
	time.sleep(2.5)
	if r == 1 :
		annoyingad()

#create new citation
	time.sleep(4.5)
	pyautogui.moveTo(15,602)
	pyautogui.click()
	pyautogui.hotkey('tab')
	pyautogui.hotkey('enter')

#cite website
	time.sleep(3)
	pyautogui.moveTo(15,602)
	pyautogui.click()
	pyautogui.hotkey('tab')
	pyautogui.hotkey('tab')
	pyautogui.hotkey('enter')


	r = r + 1



#back to citations
time.sleep(2.5)
pyautogui.moveTo(15,602)
pyautogui.click()
pyautogui.hotkey('tab')
pyautogui.hotkey('enter')


#Copying the final Citations

time.sleep(5.5)
pyautogui.moveTo(15,602)
pyautogui.click()
pyautogui.hotkey('tab')
pyautogui.hotkey('tab')
pyautogui.hotkey('tab')
pyautogui.hotkey('tab')
pyautogui.hotkey('tab')
pyautogui.hotkey('enter')


time.sleep(2)
pyautogui.hotkey('alt', 'tab')

time.sleep(1)
pyautogui.hotkey('ctrl', 'v')


#
def annoyingad():
	pyautogui.moveTo(979,568)
	pyautogui.click()
	time.sleep(20)


driver.close()